"""Build an editable, print-oriented DOCX booklet from the online program data."""
from __future__ import annotations

import json
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
REPO = ROOT.parents[1]
OUT = ROOT / "嘉義高中管樂社第41屆校友聯合音樂會_為伍_紙本節目冊.docx"
NODE = "/Users/linjiunyu/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"

# A4 print booklet override: 16 mm / 18 mm margins; Noto CJK typesetting; deep ink + brass gold.
INK = "173247"
GOLD = "A77B35"
PAPER = "FBF8F1"
MIST = "EAF0F1"
MUTED = "63727A"
# The supplied macOS font is embeddable and renders reliably in Word/LibreOffice.
FONT = "PingFang TC"
SANS = "PingFang TC"


def get_data():
    js = ROOT / "data" / "concert-41st.js"
    command = (
        "const fs=require('fs'),vm=require('vm');"
        "const s={window:{}};vm.createContext(s);"
        "vm.runInContext(fs.readFileSync(process.argv[1],'utf8'),s);"
        "process.stdout.write(JSON.stringify(s.window.CONCERT_PROGRAM_DATA));"
    )
    raw = subprocess.check_output([NODE, "-e", command, str(js)], text=True)
    return json.loads(raw)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=150, bottom=100, end=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    rep = OxmlElement("w:tblHeader")
    rep.set(qn("w:val"), "true")
    trPr.append(rep)


def set_fixed_table(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            tcW.set(qn("w:w"), str(round(width / 2.54 * 1440)))
            tcW.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(p, before=0, after=4, line=1.35, align=None, keep=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        pPr = p._p.get_or_add_pPr()
        keepNext = OxmlElement("w:keepNext")
        pPr.append(keepNext)


def add_text(doc, text, size=10.5, color=INK, bold=False, italic=False, before=0, after=4,
             line=1.35, align=None, style=None, keep=False):
    p = doc.add_paragraph(style=style)
    set_para(p, before, after, line, align, keep)
    set_run(p.add_run(text), size, color, bold, italic)
    return p


def add_label(doc, text, after=5, before=0):
    p = doc.add_paragraph()
    set_para(p, before, after, 1.0, keep=True)
    set_run(p.add_run(text.upper()), 8.5, GOLD, True, font=SANS)
    return p


def add_section_title(doc, title, en=None, before=12):
    add_label(doc, en or "PROGRAM BOOK", after=2, before=before)
    p = doc.add_paragraph()
    set_para(p, 0, 10, 1.15, keep=True)
    set_run(p.add_run(title), 20, INK, True)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(paragraph.add_run("嘉義高中管樂社｜第 41 屆《為伍》　"), 8.5, MUTED, font=SANS)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_header_footer(section):
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run("《為伍》 KEEP COMPANY  |  2026"), 8, GOLD, True, font=SANS)
    add_page_number(section.footer.paragraphs[0])


def no_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""


def add_image(doc, path, width_cm, caption=None):
    if not path.exists():
        return
    # python-docx accepts PNG/JPEG but not WebP; retain the original asset in
    # the site and make a temporary PNG only for DOCX embedding.
    image_path = path
    if path.suffix.lower() == ".webp":
        temp_dir = Path(tempfile.gettempdir()) / "cysh-41st-program-docx-images"
        temp_dir.mkdir(parents=True, exist_ok=True)
        image_path = temp_dir / f"{path.stem}.png"
        if not image_path.exists() or image_path.stat().st_mtime < path.stat().st_mtime:
            with Image.open(path) as image:
                image.convert("RGB").save(image_path, "PNG", optimize=True)
    p = doc.add_paragraph()
    set_para(p, 2, 3, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        add_text(doc, caption, 8, MUTED, italic=True, after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)


def page_break(doc):
    doc.add_page_break()


def add_program_card(doc, item):
    # Keep one entry coherent; Word can split long notes naturally after the opening block.
    table = doc.add_table(rows=1, cols=2)
    set_fixed_table(table, [1.15, 15.35])
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(left, INK)
    set_cell_shading(right, PAPER)
    set_cell_margins(left, 140, 130, 140, 130)
    set_cell_margins(right, 150, 180, 150, 180)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, 0, 0, 1.0)
    set_run(p.add_run(f"{item['no']:02d}"), 19, "FFFFFF", True, font=SANS)
    p = right.paragraphs[0]
    set_para(p, 0, 1, 1.1)
    zh = next((x["text"] for x in item["titles"] if x["lang"] == "zh-Hant"), item["titles"][0]["text"])
    set_run(p.add_run(zh), 13.5, INK, True)
    foreign = [x["text"] for x in item["titles"] if x["lang"] != "zh-Hant"]
    if foreign:
        p2 = right.add_paragraph()
        set_para(p2, 0, 4, 1.05)
        set_run(p2.add_run(" / ".join(foreign)), 9.5, GOLD, italic=True, font=SANS)
    comp = item.get("composer", "")
    comp_text = "、".join(comp) if isinstance(comp, list) else comp
    arr = item.get("arranger", "")
    arr_text = "、".join(arr) if isinstance(arr, list) else arr
    meta = f"作曲｜{comp_text}"
    if arr_text:
        meta += f"　編曲｜{arr_text}"
    if item.get("soloist"):
        meta += f"　獨奏｜{item['soloist']}"
    meta += f"　演出時間｜{item.get('duration', '')}"
    p3 = right.add_paragraph()
    set_para(p3, 0, 0, 1.2)
    set_run(p3.add_run(meta), 8.8, MUTED, font=SANS)
    for paragraph in item.get("note", []):
        add_text(doc, paragraph, 10, INK, after=5, line=1.45)
    add_text(doc, "", after=2, line=1.0)


def add_person(doc, person, label):
    photo = ROOT / person["photo"]
    if photo.exists():
        add_image(doc, photo, 12.8, person.get("photoAlt"))
    add_label(doc, label, after=2)
    p = doc.add_paragraph()
    set_para(p, 0, 2, 1.1, keep=True)
    set_run(p.add_run(person["name"]), 16, INK, True)
    p.add_run("　")
    set_run(p.add_run(f"{person.get('role', '')}　{person.get('number', '')}"), 9.5, GOLD, True, font=SANS)
    add_text(doc, person["bio"], 10.2, INK, after=10, line=1.45)


def add_ensemble(doc, ensemble):
    photo = ROOT / ensemble["photo"]
    if photo.exists():
        add_image(doc, photo, 15.7, ensemble.get("photoAlt"))
    add_label(doc, ensemble["subtitle"], after=2)
    add_text(doc, ensemble["title"], 15, INK, True, after=5, line=1.1)
    for para in ensemble["content"]:
        add_text(doc, para, 10.2, INK, after=5, line=1.45)
    add_text(doc, "", after=5, line=1.0)


def main():
    data = get_data()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    no_header_footer(section)

    # A restrained, editable A4 booklet design (rather than a screen capture).
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for style_name, size, color in (("Heading 1", 20, INK), ("Heading 2", 14, INK), ("Heading 3", 11, GOLD)):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    info = data["info"]

    # Cover
    add_text(doc, "嘉義高中管樂社", 10, GOLD, True, after=12, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_para(p, 0, 2, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run("第 41 屆"), 15, INK, True)
    p = doc.add_paragraph()
    set_para(p, 0, 2, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run("《為伍》"), 37, INK, True)
    add_text(doc, "KEEP COMPANY", 12, GOLD, True, after=15, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    poster = REPO / "assets" / "img" / "poster_weiwu_2026.webp"
    add_image(doc, poster, 13.5)
    add_text(doc, "嘉義高中校友暨在校生聯合音樂會", 12, INK, True, after=7, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"{info['date']}　{info['dayOfWeek']}　{info['time']}", 11, GOLD, True, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, info["venue"], 10.5, INK, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, info["venueAddress"], 8.5, MUTED, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(1.6), Cm(1.6)
    section.left_margin, section.right_margin = Cm(1.75), Cm(1.75)
    add_header_footer(section)

    # Concert information / opening note
    add_section_title(doc, "演出資訊", "CONCERT INFORMATION", before=0)
    table = doc.add_table(rows=0, cols=2)
    for label, value in [
        ("演出日期", f"{info['date']}（{info['dayOfWeek']}） {info['time']}"),
        ("入場時間", info["doorTime"]),
        ("演出地點", f"{info['venue']}｜{info['venueAddress']}"),
        ("主辦字頭", info["organizer"]),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = "", ""
        set_cell_shading(cells[0], INK)
        set_cell_shading(cells[1], PAPER)
        for c in cells:
            set_cell_margins(c, 90, 145, 90, 145)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = cells[0].paragraphs[0]; set_para(p0, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER); set_run(p0.add_run(label), 9, "FFFFFF", True, font=SANS)
        p1 = cells[1].paragraphs[0]; set_para(p1, 0, 0, 1.2); set_run(p1.add_run(value), 10, INK)
    set_fixed_table(table, [3.15, 13.35])
    add_section_title(doc, data["presidentMessage"]["title"], "A MESSAGE FROM THE PRESIDENT", before=16)
    president = data["presidentMessage"]
    add_image(doc, ROOT / president["photo"], 12.8, president["photoAlt"])
    add_text(doc, president["subtitle"], 12, GOLD, True, italic=True, after=7, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"{president['author']}｜{president['number']}", 9.5, MUTED, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    for para in president["content"]:
        add_text(doc, para, 10.3, INK, after=6, line=1.48)

    # Program order
    page_break(doc)
    add_section_title(doc, "演出曲目", "PROGRAM AT A GLANCE", before=0)
    for heading, songs in (("上半場", data["program"]["firstHalf"]), ("下半場", data["program"]["secondHalf"])):
        add_label(doc, heading, after=4, before=4)
        for item in songs:
            zh = next((x["text"] for x in item["titles"] if x["lang"] == "zh-Hant"), item["titles"][0]["text"])
            foreign = next((x["text"] for x in item["titles"] if x["lang"] == "en"), "")
            p = doc.add_paragraph()
            set_para(p, 0, 3, 1.2)
            set_run(p.add_run(f"{item['no']:02d}  "), 10, GOLD, True, font=SANS)
            set_run(p.add_run(zh), 11, INK, True)
            if foreign:
                set_run(p.add_run(f"  {foreign}"), 8.5, MUTED, italic=True, font=SANS)
            set_run(p.add_run(f"　{item.get('duration','')}"), 8.5, GOLD, font=SANS)
    add_image(doc, ROOT / data["program"]["heroImage"]["src"], 15.6, data["program"]["heroImage"]["alt"])

    # Program notes
    page_break(doc)
    add_section_title(doc, "曲目導賞", "PROGRAM NOTES", before=0)
    add_label(doc, "上半場", after=6)
    for item in data["program"]["firstHalf"]:
        add_program_card(doc, item)
    add_label(doc, "下半場", after=6, before=8)
    for item in data["program"]["secondHalf"]:
        add_program_card(doc, item)

    # People
    page_break(doc)
    add_section_title(doc, "團隊與獨奏", "PEOPLE ON STAGE", before=0)
    for conductor in data["leadership"]["conductors"]:
        add_person(doc, conductor, "樂團指揮")
    add_person(doc, data["leadership"]["soloist"][0], "小號獨奏")
    for ensemble in data["leadership"]["ensembles"]:
        add_ensemble(doc, ensemble)

    # Roster
    page_break(doc)
    add_section_title(doc, "演出人員名冊", "MUSICIAN ROSTER", before=0)
    for group in data["roster"]:
        add_label(doc, group["section"], after=2, before=7)
        p = doc.add_paragraph()
        set_para(p, 0, 4, 1.2)
        set_run(p.add_run(group["sectionZh"]), 12.5, INK, True)
        members = "　".join(f"{m['name']}（{m['number']}）" for m in group["members"])
        add_text(doc, members, 10.2, INK, after=4, line=1.5)

    # Staff and thanks
    page_break(doc)
    add_section_title(doc, data["organization"]["staffTitle"], "PRODUCTION TEAM", before=0)
    staff = data["organization"]["staffGroups"]
    table = doc.add_table(rows=0, cols=2)
    for i, group in enumerate(staff):
        cells = table.add_row().cells
        for c in cells:
            set_cell_margins(c, 90, 145, 90, 145)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], INK if i % 2 == 0 else "29495C")
        set_cell_shading(cells[1], PAPER)
        p0 = cells[0].paragraphs[0]; set_para(p0, 0, 0, 1.0); set_run(p0.add_run(group["role"]), 9, "FFFFFF", True, font=SANS)
        p1 = cells[1].paragraphs[0]; set_para(p1, 0, 0, 1.2); set_run(p1.add_run("、".join(group["names"])), 10, INK)
    set_fixed_table(table, [5.3, 11.2])
    add_section_title(doc, data["organization"]["thanksTitle"], "WITH GRATITUDE", before=18)
    add_image(doc, ROOT / data["organization"]["heroImage"]["src"], 15.6, data["organization"]["heroImage"]["alt"])
    for thanks in data["organization"]["thanksList"]:
        p = doc.add_paragraph()
        set_para(p, 0, 3, 1.25)
        set_run(p.add_run("◆  "), 9, GOLD, True, font=SANS)
        set_run(p.add_run(thanks), 10.5, INK)
    add_text(doc, "資料來源：嘉義高中管樂社第 41 屆《為伍》線上節目冊。此版本為可編輯之紙本排版檔。", 8, MUTED, after=0, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Document metadata
    props = doc.core_properties
    props.title = "第 41 屆嘉義高中校友暨在校生聯合音樂會《為伍》紙本節目冊"
    props.subject = "2026.08.08 嘉義市政府文化局音樂廳"
    props.author = "嘉義高中管樂社"
    props.comments = "依線上節目冊內容編排為可編輯 A4 紙本節目冊。"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
